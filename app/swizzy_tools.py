import os
import logging
import io # Import io for BytesIO
import pandas as pd # Import pandas
from app.services.storage_service import storage_service  # Import shared service instance
from pypdf import PdfReader
import docx # python-docx library
from agents import function_tool # Import the decorator
import openpyxl # Add full openpyxl import
from openpyxl import load_workbook # Needed for reading existing xlsx
from openpyxl.utils import get_column_letter, column_index_from_string, range_boundaries # <-- Import openpyxl utils
from openpyxl.styles import Font, PatternFill, Alignment # <-- Import Font
import json # <-- Import JSON
import contextlib # <-- For capturing stdout
import traceback # <-- For formatting exceptions
import numpy as np # <-- Import numpy for restricted env
import time # For unique filenames
import uuid # For unique filenames
from app.services.gemini_service import GeminiService  # Add this import

# --- Sklearn Imports (for restricted env) ---
try:
    from sklearn import preprocessing, cluster, linear_model, metrics, model_selection
    _sklearn_available = True
except ImportError:
    _sklearn_available = False
    # Handle error within the tool if sklearn not installed
    pass

# --- OCR Imports ---
try:
    import pytesseract
    from PIL import Image
    _pytesseract_available = True
except ImportError:
    _pytesseract_available = False
    # We'll handle the error message within the tool if called without library
    pass 

logger = logging.getLogger(__name__)

# Helper to convert Excel cell notation ('A1') to zero-based row/col index
def excel_coord_to_indices(coord: str) -> tuple[int, int]:
    """Converts 'A1' style coords to (row_idx, col_idx) zero-based."""
    col_str, row_str = '', ''
    for char in coord:
        if char.isalpha():
            col_str += char
        else:
            row_str += char
    
    # Convert column letters to 0-indexed number (A->0, B->1, etc.)
    col_idx = 0
    for char in col_str.upper():
        col_idx = col_idx * 26 + (ord(char) - ord('A') + 1)
    col_idx -= 1  # Convert to 0-indexed
    
    # Convert row to 0-indexed
    row_idx = int(row_str) - 1
    
    return row_idx, col_idx

@function_tool
def read_file_content(file_handle: str) -> str:
    """
    Reads the content of a file identified by its handle.

    Args:
        file_handle: The handle of the file to read.

    Returns:
        The content of the file as text.
    """
    logger.info(f"Reading file content for handle: {file_handle}")
    print(f"Reading file content for handle: {file_handle}")
    
    try:
        # Get the file bytes from storage
        file_bytes = storage_service.download_file(file_handle)
        
        # Determine file type based on extension
        file_extension = os.path.splitext(file_handle)[1].lower()
        
        if file_extension in ['.txt', '.md']:
            # For text and markdown files, decode as UTF-8
            text_content = file_bytes.decode('utf-8', errors='replace')
            logger.info(f"Successfully read {file_extension} content from handle: {file_handle}")
        
        elif file_extension == '.pdf':
            # For PDF files, use PyPDF to extract text
            pdf_io = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_io)
            
            text_content = ""
            for page in reader.pages:
                text_content += page.extract_text() + "\n\n"
            
            logger.info(f"Successfully read PDF content from handle: {file_handle}")
        
        elif file_extension == '.docx':
            # For DOCX files, use python-docx
            docx_io = io.BytesIO(file_bytes)
            doc = docx.Document(docx_io)
            
            text_content = ""
            for para in doc.paragraphs:
                text_content += para.text + "\n"
            
            logger.info(f"Successfully read DOCX content from handle: {file_handle}")
        
        elif file_extension in ['.xlsx', '.csv']:
            # For Excel/CSV files
            try:
                if file_extension == '.csv':
                    # Use pandas to read CSV
                    csv_io = io.StringIO(file_bytes.decode('utf-8', errors='replace'))
                    df = pd.read_csv(csv_io)
                    text_content = df.to_string(index=False)
                    logger.info(f"Successfully read CSV content from handle: {file_handle}")
                
                else:  # .xlsx
                    # Use openpyxl to read Excel
                    excel_io = io.BytesIO(file_bytes)
                    wb = load_workbook(excel_io, read_only=True)
                    
                    combined_text = []
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                        sheet_text = f"Sheet: {sheet_name}\n"
                        
                        for row in sheet.iter_rows(values_only=True):
                            row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                            sheet_text += row_text + "\n"
                        
                        combined_text.append(sheet_text)
                
                text_content = "\n\n".join(combined_text)  # Join sheets with double newlines
                logger.info(f"Successfully read XLSX content from handle: {file_handle}")
            except Exception as e:
                logger.error(f"Error reading Excel file '{file_handle}': {e}")
                return f"Error processing Excel file '{file_handle}': {e}"

        else:
            logger.warning(f"Unsupported file type '{file_extension}' for handle: {file_handle}")
            return f"Unsupported file type: {file_extension}. Please provide a .txt, .md, .pdf, .docx, .xlsx, or .csv file."
        
        return text_content
    
    except FileNotFoundError:
        logger.error(f"File not found for handle: {file_handle}")
        return f"Error: File not found for handle '{file_handle}'"
    
    except Exception as e:
        logger.exception(f"Error reading file content for handle '{file_handle}': {e}")
        return f"Error processing file '{file_handle}': {e}"

@function_tool
def ponder_document_request(request_description: str, points_to_consider: str) -> str:
    """
    A tool for the document agent to think through a request before taking action.
    This allows for more careful planning and consideration of the appropriate approach.
    
    Args:
        request_description: A brief description of the client's request
        points_to_consider: Key points to consider when handling this request
        
    Returns:
        Guidance on how to proceed with the request
    """
    print(f"Document agent pondering request: {request_description}")
    print(f"Points to consider: {points_to_consider}")
    
    # For now, this just echoes back the points with some generic guidance
    response = {
        "analysis": points_to_consider,
        "recommendation": "Based on your analysis, proceed with the appropriate document tool calls to fulfill the request. Remember to use read_file_content first if working with an existing document, create_document for new documents, or extract_text_from_image if there are images involved."
    }
    
    return json.dumps(response)

@function_tool
def create_document(filename: str, content: str) -> str:
    """
    Creates a new document file (e.g., .docx or .txt) with the given content.

    Args:
        filename: The desired name for the document file (including extension).
        content: The content to put into the document.

    Returns:
        A message indicating success or failure, potentially including the file handle of the created file.
    """
    logger.info(f"Attempting to create document '{filename}'.")
    
    # Ensure filename has a valid extension, default to .txt if not
    if not filename.lower().endswith(('.docx', '.txt')):
        logger.warning(f"Filename '{filename}' lacks document extension. Appending '.txt'.")
        filename += '.txt'
    
    try:
        if filename.lower().endswith('.txt'):
            # For .txt files, just write the content directly
            file_bytes = content.encode('utf-8')
            
        elif filename.lower().endswith('.docx'):
            # For .docx files, use python-docx to create a document
            doc = docx.Document()
            
            # Split content by newlines and add each paragraph
            for paragraph in content.split('\n'):
                if paragraph.strip():  # Skip empty paragraphs
                    doc.add_paragraph(paragraph)
            
            # Save to a BytesIO object
            docx_io = io.BytesIO()
            doc.save(docx_io)
            file_bytes = docx_io.getvalue()
        
        # Upload the file bytes using the storage service
        file_handle = storage_service.upload_file(filename, file_bytes)
        logger.info(f"Successfully created document '{filename}' with handle '{file_handle}'.")
        return f"Successfully created document '{filename}'. Handle: {file_handle}"
    
    except Exception as e:
        logger.exception(f"Error creating document '{filename}': {e}")
        return f"Error creating document '{filename}': {e}"

@function_tool
def extract_text_from_image(file_handle: str) -> str:
    """
    Extracts text from an image file using OCR.

    Args:
        file_handle: The handle of the image file to process.

    Returns:
        The extracted text from the image.
    """
    logger.info(f"Attempting to extract text from image '{file_handle}'.")
    
    if not _pytesseract_available:
        return "Error: OCR functionality is not available. The pytesseract library is not installed."
    
    try:
        # Get the file bytes from storage
        file_bytes = storage_service.download_file(file_handle)
        
        # Load the image using PIL
        image_io = io.BytesIO(file_bytes)
        image = Image.open(image_io)
        
        # Use pytesseract to extract text
        extracted_text = pytesseract.image_to_string(image)
        
        if not extracted_text.strip():
            logger.warning(f"No text extracted from image '{file_handle}'.")
            return "No text could be extracted from the image. The image may not contain text or the text may be unclear."
        
        logger.info(f"Successfully extracted text from image '{file_handle}'.")
        return extracted_text
    
    except FileNotFoundError:
        logger.error(f"Image file not found for handle: {file_handle}")
        return f"Error: Image file not found for handle '{file_handle}'"
    
    except Exception as e:
        logger.exception(f"Error during text extraction from image '{file_handle}': {e}")
        return f"Error processing image '{file_handle}': {e}"

def _clean_csv_block(block: str) -> str:
    """Helper to clean and validate a CSV block."""
    lines = []
    for line in block.splitlines():
        line = line.strip()
        if line.startswith('```') or not line:
            continue
        # Only process lines that are actual CSV content
        if ',' in line:
            lines.append(line)
    return '\n'.join(lines)

@function_tool
def extract_spreadsheet_from_document(file_handle: str, extraction_instruction: str = "", output_format: str = "csv") -> str:
    """
    Extracts spreadsheet data from documents using AI. Handles PDFs, images, and text files.
    Automatically converts to Excel for multiple tables.
    """
    logger.info(f"Extracting spreadsheet data from: {file_handle}")
    
    try:
        file_bytes = storage_service.download_file(file_handle)
        file_extension = os.path.splitext(file_handle)[1].lower()
        gemini_service = GeminiService()
        extracted_data = []

        # Base prompt for all file types
        base_prompt = f"""Extract all tables from this document and return them as CSV data.
        Format requirements:
        - Use ```csv blocks for each table
        - Include headers
        - Maintain data relationships
        - Clean formatting
        Specific instructions: {extraction_instruction}
        """

        # Get AI response based on file type
        if file_extension == '.pdf':
            result = gemini_service.analyze_pdf(file_bytes, base_prompt)
        elif file_extension in ['.jpg', '.jpeg', '.png']:
            result = gemini_service.analyze_image(file_bytes, base_prompt)
        else:  # Text or markdown
            text_content = file_bytes.decode('utf-8', errors='replace')
            result = gemini_service.analyze_text(text_content, base_prompt)

        if not result.get("success"):
            return f"Error: AI extraction failed: {result.get('error', 'Unknown error')}"

        # Extract CSV blocks from AI response
        blocks = result["text"].split('```csv')
        for block in blocks[1:]:  # Skip first empty split
            if block.strip():
                end_marker = block.find('```')
                if end_marker > -1:
                    csv_content = block[:end_marker].strip()
                    if csv_content:
                        extracted_data.append(csv_content)

        if not extracted_data:
            return "Error: No valid tables found in document"

        # Auto-select output format based on number of tables
        if len(extracted_data) > 1:
            output_format = "xlsx"
            logger.info(f"Multiple tables detected ({len(extracted_data)}), using Excel format")
        
        timestamp = int(time.time())
        base_name = os.path.splitext(os.path.basename(file_handle))[0]
        output_filename = f"{base_name}_extracted_{timestamp}.{output_format}"

        if output_format == "csv":
            file_bytes = extracted_data[0].encode('utf-8')  # Just use first table
            file_handle = storage_service.upload_file(output_filename, file_bytes)
            return f"Extracted single table to CSV. Handle: {file_handle}"
        
        else:  # xlsx for multiple tables
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                for i, csv_block in enumerate(extracted_data, 1):
                    try:
                        df = pd.read_csv(io.StringIO(csv_block))
                        sheet_name = f'Table_{i}'
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                        writer.sheets[sheet_name].sheet_state = 'visible'
                    except Exception as e:
                        logger.warning(f"Error processing table {i}: {e}")
                        continue

            file_handle = storage_service.upload_file(output_filename, buffer.getvalue())
            return f"Extracted {len(extracted_data)} tables to Excel. Handle: {file_handle}"

    except Exception as e:
        logger.exception(f"Error extracting spreadsheet data: {e}")
        return f"Error processing document: {str(e)}"

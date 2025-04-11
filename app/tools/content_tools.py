import os
import logging
import io
import json
import time
import re
from typing import Dict, List, Any, Optional, Union, Tuple
from app.services import storage_service
from app.services.gemini_service import GeminiService
from agents import function_tool

# Add imports for high-fidelity DOCX conversion
import docx
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

# Add imports for PDF handling
from pypdf import PdfReader
import fitz  # PyMuPDF
import base64
from PIL import Image

logger = logging.getLogger(__name__)

def convert_docx_to_markdown(docx_bytes: bytes) -> str:
    """
    Convert DOCX content to Markdown format with high fidelity.
    
    Args:
        docx_bytes: The bytes of the DOCX file
        
    Returns:
        Markdown content as a string
    """
    try:
        # Load the DOCX document
        docx_io = io.BytesIO(docx_bytes)
        doc = docx.Document(docx_io)
        
        # Extract document properties
        properties = {}
        core_properties = doc.core_properties
        if core_properties:
            if core_properties.title:
                properties["title"] = core_properties.title
            if core_properties.author:
                properties["author"] = core_properties.author
            if core_properties.created:
                properties["created"] = core_properties.created.isoformat() if hasattr(core_properties.created, 'isoformat') else str(core_properties.created)
            if core_properties.modified:
                properties["modified"] = core_properties.modified.isoformat() if hasattr(core_properties.modified, 'isoformat') else str(core_properties.modified)
        
        # Start building the markdown content
        md_lines = []
        
        # Add YAML frontmatter with document properties
        if properties:
            md_lines.append("---")
            for key, value in properties.items():
                md_lines.append(f"{key}: {value}")
            md_lines.append("---")
            md_lines.append("")
        
        # Process document elements
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Process paragraph
                paragraph = Paragraph(element, doc)
                md_lines.append(process_paragraph(paragraph))
            elif isinstance(element, CT_Tbl):
                # Process table
                table = Table(element, doc)
                md_lines.append(process_table(table))
        
        # Join the markdown lines
        markdown_content = "\n".join(md_lines)
        
        return markdown_content
    except Exception as e:
        logger.exception(f"Error converting DOCX to Markdown: {e}")
        return f"Error converting DOCX to Markdown: {e}"

def process_paragraph(paragraph: Paragraph) -> str:
    """
    Process a DOCX paragraph and convert it to Markdown.
    
    Args:
        paragraph: The DOCX paragraph
        
    Returns:
        Markdown representation of the paragraph
    """
    # Check if it's a heading
    if paragraph.style.name.startswith('Heading'):
        level = int(paragraph.style.name.replace('Heading', ''))
        return f"{'#' * level} {paragraph.text}"
    
    # Check if it's a list item
    if paragraph.style.name.startswith('List'):
        return f"- {paragraph.text}"
    
    # Process text with formatting
    md_text = ""
    for run in paragraph.runs:
        text = run.text
        
        # Apply formatting
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        if run.underline:
            text = f"__{text}__"
        if run.strike:
            text = f"~~{text}~~"
        
        md_text += text
    
    # Handle hyperlinks
    for hyperlink in paragraph._element.xpath(".//w:hyperlink", namespaces=paragraph._element.nsmap):
        for run_element in hyperlink.xpath(".//w:r", namespaces=paragraph._element.nsmap):
            run_text = "".join([t.text for t in run_element.xpath(".//w:t", namespaces=paragraph._element.nsmap)])
            
            # Get the relationship ID
            rel_id = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if rel_id:
                # Get the target URL from the relationship
                target_url = ""
                for rel in paragraph.part.rels:
                    if rel == rel_id:
                        target_url = paragraph.part.rels[rel].target_ref
                        break
                
                if target_url:
                    # Replace the text with a markdown link
                    md_text = md_text.replace(run_text, f"[{run_text}]({target_url})")
    
    return md_text

def process_table(table: Table) -> str:
    """
    Process a DOCX table and convert it to Markdown.
    
    Args:
        table: The DOCX table
        
    Returns:
        Markdown representation of the table
    """
    md_lines = []
    
    # Process each row
    for i, row in enumerate(table.rows):
        # Get cell text for this row
        cells = [cell.text.replace('\n', ' ') for cell in row.cells]
        md_lines.append(f"| {' | '.join(cells)} |")
        
        # Add header separator after first row
        if i == 0:
            md_lines.append(f"|{' --- |' * len(cells)}")
    
    # Add an empty line after the table
    md_lines.append("")
    
    return "\n".join(md_lines)

def convert_markdown_to_docx(markdown_content: str) -> bytes:
    """
    Convert Markdown content to DOCX format with high fidelity.
    
    Args:
        markdown_content: The Markdown content as a string
        
    Returns:
        DOCX content as bytes
    """
    try:
        # Create a new DOCX document
        doc = docx.Document()
        
        # Extract YAML frontmatter if present
        frontmatter = {}
        if markdown_content.startswith('---'):
            end_index = markdown_content.find('---', 3)
            if end_index > 0:
                frontmatter_text = markdown_content[3:end_index].strip()
                for line in frontmatter_text.split('\n'):
                    if ':' in line:
                        key, value = line.split(':', 1)
                        frontmatter[key.strip()] = value.strip()
                
                # Remove frontmatter from content
                markdown_content = markdown_content[end_index+3:].strip()
        
        # Set document properties from frontmatter
        if 'title' in frontmatter:
            doc.core_properties.title = frontmatter['title']
        if 'author' in frontmatter:
            doc.core_properties.author = frontmatter['author']
        
        # Process the markdown content line by line
        lines = markdown_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Process headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                doc.add_heading(text, level=level)
                i += 1
                continue
            
            # Process tables
            if line.startswith('|') and i+1 < len(lines) and lines[i+1].startswith('|') and '---' in lines[i+1]:
                table_lines = []
                while i < len(lines) and lines[i].startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                
                # Parse table
                rows = []
                for j, table_line in enumerate(table_lines):
                    if j == 1:  # Skip separator line
                        continue
                    
                    # Extract cells from the line
                    cells = table_line.strip('|').split('|')
                    rows.append([cell.strip() for cell in cells])
                
                # Create table in document
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    
                    # Fill table cells
                    for row_idx, row in enumerate(rows):
                        for col_idx, cell_text in enumerate(row):
                            if col_idx < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[col_idx].text = cell_text
                
                continue
            
            # Process lists
            list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
            if list_match:
                indent = len(list_match.group(1))
                text = list_match.group(3)
                
                # Add list paragraph with appropriate style
                p = doc.add_paragraph(text)
                p.style = 'List Bullet' if list_match.group(2) in ['-', '*', '+'] else 'List Number'
                
                i += 1
                continue
            
            # Process regular paragraphs with formatting
            if line.strip():
                # Handle basic formatting
                line = re.sub(r'\*\*(.+?)\*\*', lambda m: m.group(1), line)  # Bold
                line = re.sub(r'\*(.+?)\*', lambda m: m.group(1), line)      # Italic
                line = re.sub(r'__(.+?)__', lambda m: m.group(1), line)      # Underline
                line = re.sub(r'~~(.+?)~~', lambda m: m.group(1), line)      # Strikethrough
                
                # Handle links
                line = re.sub(r'\[(.+?)\]\((.+?)\)', lambda m: m.group(1), line)
                
                p = doc.add_paragraph(line)
                
                # TODO: Apply actual formatting to runs
                # This would require parsing the markdown formatting and applying
                # it to the appropriate runs in the paragraph
            
            i += 1
        
        # Save to bytes
        docx_io = io.BytesIO()
        doc.save(docx_io)
        return docx_io.getvalue()
    except Exception as e:
        logger.exception(f"Error converting Markdown to DOCX: {e}")
        raise ValueError(f"Error converting Markdown to DOCX: {e}")

def extract_pdf_content_with_formatting(pdf_bytes: bytes) -> str:
    """
    Extract content from a PDF file with formatting preserved as Markdown.
    
    Args:
        pdf_bytes: The bytes of the PDF file
        
    Returns:
        Markdown content as a string with tables and image references
    """
    try:
        # Load the PDF using PyMuPDF (fitz)
        pdf_io = io.BytesIO(pdf_bytes)
        pdf_document = fitz.open(stream=pdf_io, filetype="pdf")
        
        # Extract document properties
        properties = {}
        if pdf_document.metadata:
            if pdf_document.metadata.get("title"):
                properties["title"] = pdf_document.metadata.get("title")
            if pdf_document.metadata.get("author"):
                properties["author"] = pdf_document.metadata.get("author")
            if pdf_document.metadata.get("creationDate"):
                properties["created"] = pdf_document.metadata.get("creationDate")
            if pdf_document.metadata.get("modDate"):
                properties["modified"] = pdf_document.metadata.get("modDate")
        
        # Start building the markdown content
        md_lines = []
        
        # Add YAML frontmatter with document properties
        if properties:
            md_lines.append("---")
            for key, value in properties.items():
                md_lines.append(f"{key}: {value}")
            md_lines.append("---")
            md_lines.append("")
        
        # Extract images to a temporary directory
        image_count = 0
        image_references = []
        
        # Process each page
        for page_num, page in enumerate(pdf_document):
            # Add page header
            md_lines.append(f"## Page {page_num + 1}")
            md_lines.append("")
            
            # Extract text blocks with formatting
            blocks = page.get_text("dict")["blocks"]
            
            # Track current heading level
            current_heading_level = 2  # Start after the page header
            
            # Process blocks
            for block in blocks:
                # Handle text blocks
                if block["type"] == 0:  # Text block
                    for line in block["lines"]:
                        line_text = ""
                        
                        # Get font information from first span
                        if line["spans"]:
                            first_span = line["spans"][0]
                            font_size = first_span["size"]
                            is_bold = "bold" in first_span["font"].lower()
                            
                            # Determine if this is a heading based on font size
                            if font_size > 14:
                                # Larger text is likely a heading
                                heading_level = min(current_heading_level, 3)  # Don't go deeper than h3
                                line_text = f"{'#' * heading_level} "
                            
                        # Process spans in the line
                        for span in line["spans"]:
                            span_text = span["text"]
                            
                            # Apply formatting based on font properties
                            if "bold" in span["font"].lower():
                                span_text = f"**{span_text}**"
                            if "italic" in span["font"].lower():
                                span_text = f"*{span_text}*"
                            
                            line_text += span_text
                        
                        # Add the processed line
                        if line_text.strip():
                            md_lines.append(line_text)
                
                # Handle image blocks
                elif block["type"] == 1:  # Image block
                    image_count += 1
                    image_name = f"image_{page_num + 1}_{image_count}.png"
                    
                    # Extract image
                    try:
                        xref = block["xref"]
                        image = pdf_document.extract_image(xref)
                        if image:
                            # For now, we'll just reference the image
                            # In a real implementation, we would save the image and provide a link
                            md_lines.append(f"![Image {image_count} on page {page_num + 1}]({image_name})")
                            md_lines.append("")
                            
                            # Store image reference
                            image_references.append({
                                "name": image_name,
                                "page": page_num + 1,
                                "xref": xref
                            })
                    except Exception as e:
                        logger.warning(f"Failed to extract image: {e}")
                        md_lines.append(f"[Image {image_count} on page {page_num + 1}]")
                        md_lines.append("")
            
            # Extract tables from the page
            tables = page.find_tables()
            if tables:
                for i, table in enumerate(tables):
                    md_lines.append(f"### Table {i+1}")
                    md_lines.append("")
                    
                    # Convert table to Markdown format
                    md_table = []
                    
                    # Add header row
                    if table.header_rows:
                        header = table.header_rows[0]
                        header_cells = [cell.text.strip() for cell in header.cells]
                        md_table.append(f"| {' | '.join(header_cells)} |")
                        md_table.append(f"|{' --- |' * len(header_cells)}")
                    
                    # Add data rows
                    for row in table.rows:
                        if row in table.header_rows:
                            continue  # Skip header row
                        row_cells = [cell.text.strip() for cell in row.cells]
                        md_table.append(f"| {' | '.join(row_cells)} |")
                    
                    md_lines.extend(md_table)
                    md_lines.append("")
            
            # Add page separator
            md_lines.append("\n---\n")
        
        # Add image reference section if there are images
        if image_references:
            md_lines.append("## Image References")
            md_lines.append("")
            for ref in image_references:
                md_lines.append(f"- {ref['name']}: Page {ref['page']}")
            md_lines.append("")
        
        # Join the markdown lines
        markdown_content = "\n".join(md_lines)
        
        return markdown_content
    except Exception as e:
        logger.exception(f"Error extracting PDF content with formatting: {e}")
        return f"Error extracting PDF content: {e}"

@function_tool
def convert_pdf_to_markdown(file_handle: str, output_filename: str = "") -> str:
    """
    Converts a PDF file to Markdown format, preserving structure, tables, and image references.
    This allows for easier editing and manipulation of the content.
    
    Args:
        file_handle: The handle of the PDF file to convert
        output_filename: Optional name for the output Markdown file. If not provided, 
                        will use the original filename with .md extension
        
    Returns:
        The file handle of the created Markdown file
    """
    logger.info(f"Converting PDF to Markdown: {file_handle}")
    
    try:
        # Get the file bytes
        file_bytes = storage_service.download_file(file_handle)
        
        # Determine file type based on extension
        file_extension = os.path.splitext(file_handle)[1].lower()
        
        if file_extension != '.pdf':
            return f"Error: File is not a PDF. Please provide a .pdf file."
        
        # Extract content with formatting
        markdown_content = extract_pdf_content_with_formatting(file_bytes)
        
        # Determine output filename
        if not output_filename:
            base_name = os.path.splitext(os.path.basename(file_handle))[0]
            output_filename = f"{base_name}.md"
        elif not output_filename.lower().endswith('.md'):
            output_filename += '.md'
        
        # Save the markdown content to a new file
        file_bytes = markdown_content.encode('utf-8')
        new_file_handle = storage_service.upload_file(output_filename, file_bytes)
        
        logger.info(f"Successfully converted PDF to Markdown: {new_file_handle}")
        return new_file_handle
        
    except FileNotFoundError:
        logger.error(f"PDF file not found: {file_handle}")
        return f"Error: PDF file not found: {file_handle}"
    except Exception as e:
        logger.exception(f"Error converting PDF to Markdown {file_handle}: {e}")
        return f"Error converting PDF to Markdown: {e}"

@function_tool
def convert_to_markdown(file_handle: str, output_filename: str = "") -> str:
    """
    Converts various file formats (PDF, DOCX) to Markdown for easier editing.
    
    Args:
        file_handle: The handle of the file to convert
        output_filename: Optional name for the output Markdown file. If not provided, 
                        will use the original filename with .md extension
        
    Returns:
        The file handle of the created Markdown file
    """
    logger.info(f"Converting file to Markdown: {file_handle}")
    
    try:
        # Get the file bytes
        file_bytes = storage_service.download_file(file_handle)
        
        # Determine file type based on extension
        file_extension = os.path.splitext(file_handle)[1].lower()
        
        # Determine output filename if not provided
        if not output_filename:
            base_name = os.path.splitext(os.path.basename(file_handle))[0]
            output_filename = f"{base_name}.md"
        elif not output_filename.lower().endswith('.md'):
            output_filename += '.md'
        
        # Convert based on file type
        if file_extension == '.pdf':
            markdown_content = extract_pdf_content_with_formatting(file_bytes)
        elif file_extension == '.docx':
            markdown_content = convert_docx_to_markdown(file_bytes)
        else:
            return f"Error: Unsupported file type for conversion: {file_extension}. Please provide a .pdf or .docx file."
        
        # Save the markdown content to a new file
        file_bytes = markdown_content.encode('utf-8')
        new_file_handle = storage_service.upload_file(output_filename, file_bytes)
        
        logger.info(f"Successfully converted {file_extension} to Markdown: {new_file_handle}")
        return new_file_handle
        
    except FileNotFoundError:
        logger.error(f"File not found: {file_handle}")
        return f"Error: File not found: {file_handle}"
    except Exception as e:
        logger.exception(f"Error converting file to Markdown {file_handle}: {e}")
        return f"Error converting file to Markdown: {e}"

@function_tool
def read_markdown(file_handle: str) -> str:
    """
    Reads a markdown file and returns its content.
    
    Args:
        file_handle: The handle of the file to read
        
    Returns:
        The content of the markdown file
    """
    logger.info(f"Reading markdown file: {file_handle}")
    
    try:
        # Get the file bytes
        file_bytes = storage_service.download_file(file_handle)
        
        # Determine file type based on extension
        file_extension = os.path.splitext(file_handle)[1].lower()
        
        if file_extension in ['.md', '.txt']:
            # For markdown/text files, decode as UTF-8
            content = file_bytes.decode('utf-8', errors='replace')
            return content
        else:
            return f"Error: Unsupported file type: {file_extension}. Please provide a .md or .txt file, or use convert_to_markdown first."
        
    except FileNotFoundError:
        logger.error(f"File not found: {file_handle}")
        return f"Error: File not found: {file_handle}"
    except Exception as e:
        logger.exception(f"Error reading file {file_handle}: {e}")
        return f"Error reading file: {e}"

def find_section_boundaries(content: str, section_identifier: str) -> Tuple[int, int]:
    """
    Find the start and end line numbers of a section in Markdown content.
    
    Args:
        content: The Markdown content
        section_identifier: Heading text or line number range (e.g., "# Introduction" or "10-20")
        
    Returns:
        Tuple of (start_line, end_line)
    """
    lines = content.split('\n')
    
    # Check if it's a line number range
    line_range_match = re.match(r'(\d+)-(\d+)', section_identifier)
    if line_range_match:
        start_line = int(line_range_match.group(1))
        end_line = int(line_range_match.group(2))
        return max(0, start_line - 1), min(len(lines) - 1, end_line - 1)
    
    # Check if it's a heading
    heading_match = re.match(r'(#+)\s+(.*)', section_identifier)
    if heading_match:
        heading_level = len(heading_match.group(1))
        heading_text = heading_match.group(2).strip()
        
        for i, line in enumerate(lines):
            if re.match(rf'#{{{heading_level}}}\s+{re.escape(heading_text)}', line):
                # Found the heading, now find the end of the section
                for j in range(i + 1, len(lines)):
                    # Section ends at the next heading of same or higher level
                    if re.match(rf'#{{{1,heading_level}}}\s+', lines[j]):
                        return i, j - 1
                # If no next heading found, section extends to the end
                return i, len(lines) - 1
    
    # If section not found, return empty range
    return -1, -1

@function_tool
def edit_markdown_section(file_handle: str, edit_operation: str) -> str:
    """
    Performs targeted edits on a markdown file without rewriting the entire document.
    
    Args:
        file_handle: The handle of the markdown file
        edit_operation: JSON string describing the edit operation:
            {
                "operation": "insert"|"remove"|"replace",
                "target": "# Section Heading" or "10-20" (line numbers),
                "content": "New content to insert or use as replacement" (not needed for remove)
            }
        
    Returns:
        The file handle of the edited file
    """
    logger.info(f"Editing markdown file: {file_handle}")
    
    try:
        # Parse the edit operation
        try:
            op = json.loads(edit_operation)
            operation_type = op.get("operation")
            target = op.get("target")
            content = op.get("content", "")
            
            if not operation_type or not target:
                return "Error: Missing required fields in edit operation"
            
            if operation_type not in ["insert", "remove", "replace"]:
                return f"Error: Unsupported operation type: {operation_type}"
            
            if operation_type in ["insert", "replace"] and not content:
                return "Error: Content is required for insert and replace operations"
                
        except json.JSONDecodeError:
            return "Error: Invalid JSON in edit operation"
        
        # Get the file content
        file_bytes = storage_service.download_file(file_handle)
        file_extension = os.path.splitext(file_handle)[1].lower()
        
        if file_extension not in ['.md', '.txt']:
            return f"Error: Unsupported file type: {file_extension}. Please provide a .md or .txt file."
        
        original_content = file_bytes.decode('utf-8', errors='replace')
        lines = original_content.split('\n')
        
        # Find the target section
        start_line, end_line = find_section_boundaries(original_content, target)
        if start_line == -1:
            return f"Error: Target section not found: {target}"
        
        # Perform the edit operation
        if operation_type == "insert":
            # Insert content at the beginning of the target section
            new_lines = lines[:start_line] + content.split('\n') + lines[start_line:]
        elif operation_type == "remove":
            # Remove the target section
            new_lines = lines[:start_line] + lines[end_line+1:]
        elif operation_type == "replace":
            # Replace the target section
            new_lines = lines[:start_line] + content.split('\n') + lines[end_line+1:]
        
        # Join the lines back into a string
        new_content = '\n'.join(new_lines)
        
        # Upload the edited file
        file_bytes = new_content.encode('utf-8')
        storage_service.upload_file(file_handle, file_bytes)
        
        logger.info(f"Successfully edited markdown file: {file_handle}")
        return file_handle
        
    except FileNotFoundError:
        logger.error(f"File not found: {file_handle}")
        return f"Error: File not found: {file_handle}"
    except Exception as e:
        logger.exception(f"Error editing markdown file {file_handle}: {e}")
        return f"Error editing markdown file: {e}"

@function_tool
def convert_file_format(file_handle: str, target_format: str) -> str:
    """
    Converts a file from one format to another.
    
    Args:
        file_handle: The handle of the file to convert
        target_format: The target format (md or docx)
        
    Returns:
        The file handle of the converted file
    """
    logger.info(f"Converting file {file_handle} to {target_format}")
    
    try:
        # Get the file bytes
        file_bytes = storage_service.download_file(file_handle)
        
        # Determine file type based on extension
        file_extension = os.path.splitext(file_handle)[1].lower()
        filename_base = os.path.splitext(file_handle)[0]
        
        if target_format.lower() == "md":
            # Convert to Markdown
            if file_extension == '.docx':
                content = convert_docx_to_markdown(file_bytes)
                new_filename = f"{filename_base}.md"
                new_file_bytes = content.encode('utf-8')
            else:
                return f"Error: Unsupported conversion from {file_extension} to .md"
        
        elif target_format.lower() == "docx":
            # Convert to DOCX
            if file_extension in ['.md', '.txt']:
                content = file_bytes.decode('utf-8', errors='replace')
                new_file_bytes = convert_markdown_to_docx(content)
                new_filename = f"{filename_base}.docx"
            else:
                return f"Error: Unsupported conversion from {file_extension} to .docx"
        
        else:
            return f"Error: Unsupported target format: {target_format}"
        
        # Upload the converted file
        new_file_handle = storage_service.upload_file(new_filename, new_file_bytes)
        logger.info(f"Successfully converted file to {new_filename}")
        return new_file_handle
        
    except FileNotFoundError:
        logger.error(f"File not found: {file_handle}")
        return f"Error: File not found: {file_handle}"
    except Exception as e:
        logger.exception(f"Error converting file {file_handle}: {e}")
        return f"Error converting file: {e}"

@function_tool
def analyze_content_structure(file_handle: str) -> str:
    """
    Analyzes the structure of a markdown file and returns an outline.
    
    Args:
        file_handle: The handle of the markdown file
        
    Returns:
        JSON string containing the document structure
    """
    logger.info(f"Analyzing content structure of file: {file_handle}")
    
    try:
        # Get the file content
        file_bytes = storage_service.download_file(file_handle)
        file_extension = os.path.splitext(file_handle)[1].lower()
        
        if file_extension not in ['.md', '.txt']:
            return f"Error: Unsupported file type: {file_extension}. Please provide a .md or .txt file."
        
        content = file_bytes.decode('utf-8', errors='replace')
        lines = content.split('\n')
        
        # Extract headings and structure
        structure = []
        current_section = None
        section_content = []
        
        for i, line in enumerate(lines):
            # Check if line is a heading
            heading_match = re.match(r'(#+)\s+(.*)', line)
            if heading_match:
                # If we were building a section, add it to the structure
                if current_section:
                    current_section["content_length"] = len('\n'.join(section_content))
                    current_section["line_count"] = len(section_content)
                    structure.append(current_section)
                
                # Start a new section
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                current_section = {
                    "level": level,
                    "heading": heading_text,
                    "start_line": i + 1,
                    "content_preview": ""
                }
                section_content = []
            elif current_section:
                section_content.append(line)
        
        # Add the last section if there is one
        if current_section:
            current_section["content_length"] = len('\n'.join(section_content))
            current_section["line_count"] = len(section_content)
            if section_content:
                preview_text = ' '.join(section_content[:3])
                current_section["content_preview"] = (preview_text[:100] + '...') if len(preview_text) > 100 else preview_text
            structure.append(current_section)
        
        # Create the document analysis
        analysis = {
            "filename": os.path.basename(file_handle),
            "total_lines": len(lines),
            "total_words": len(content.split()),
            "structure": structure,
            "section_count": len(structure)
        }
        
        return json.dumps(analysis, indent=2)
        
    except FileNotFoundError:
        logger.error(f"File not found: {file_handle}")
        return json.dumps({"error": f"File not found: {file_handle}"})
    except Exception as e:
        logger.exception(f"Error analyzing content structure of file {file_handle}: {e}")
        return json.dumps({"error": f"Error analyzing content structure: {e}"})

@function_tool
def extract_text_from_image(file_handle: str, extraction_type: str = "text") -> str:
    """
    Extracts text or structured data from an image using Google Gemini Multi-Modal AI.
    
    Args:
        file_handle: The handle of the image file
        extraction_type: Type of extraction to perform - 'text' for general text, 
                        'invoice', 'receipt', 'table', etc. for structured data extraction
        
    Returns:
        Extracted text or structured data in JSON format
    """
    logger.info(f"Extracting {extraction_type} from image: {file_handle}")
    
    try:
        # Get the file bytes
        file_bytes = storage_service.download_file(file_handle)
        
        # Initialize Gemini service
        gemini_service = GeminiService()
        
        # Perform extraction based on type
        if extraction_type == "text":
            # Extract general text
            result = gemini_service.analyze_image(file_bytes, "Extract all text from this image, preserving the layout as much as possible.")
        else:
            # Extract structured data
            result = gemini_service.extract_structured_data(file_bytes, extraction_type)
        
        # Check if extraction was successful
        if result.get("success"):
            if extraction_type == "text":
                return result["text"]
            else:
                # Return structured data as JSON string
                return json.dumps(result["structured_data"], indent=2)
        else:
            error_message = result.get("error", "Unknown error")
            logger.error(f"Error extracting {extraction_type} from image: {error_message}")
            return f"Error extracting {extraction_type} from image: {error_message}"
            
    except FileNotFoundError:
        logger.error(f"Image file not found: {file_handle}")
        return f"Error: Image file not found: {file_handle}"
    except Exception as e:
        logger.exception(f"Error extracting {extraction_type} from image {file_handle}: {e}")
        return f"Error extracting {extraction_type} from image: {e}"


@function_tool
def create_markdown(text: str, filename: str) -> str:
    """
    Creates a markdown file with the input text as a heading 1.

    Args:
        text: The text to be used as the heading.
        filename: The name of the file to create (e.g., "my_file.md").

    Returns:
        The path to the created markdown file.
    """
    # Ensure the filename has a .md extension
    if not filename.endswith(".md"):
        filename += ".md"

    filepath = os.path.join("file_store", filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"# {text}")

    return filepath















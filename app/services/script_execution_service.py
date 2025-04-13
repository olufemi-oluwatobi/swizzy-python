import pandas as pd
import numpy as np
from typing import Dict, Any
import json
from ..services.storage_service import storage_service
from ..services.gemini_service import GeminiService
import traceback
from PIL import Image
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
# Remove circular import
# from ..tools.spreadsheet_tools import create_spreadsheet, modify_spreadsheet
from app.tools.spreadsheet_utils import read_excel_all, write_excel_file, get_excel_bytes, get_excel_bytes_from_dfs
import base64
import docx
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter, column_index_from_string
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side,
    NamedStyle, Color, Protection, Fill
)
from openpyxl.styles.colors import COLOR_INDEX
from openpyxl.worksheet.dimensions import ColumnDimension, RowDimension
from openpyxl.cell import Cell
from openpyxl.drawing.image import Image as XLImage
import PyPDF2
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from app.tools.spreadsheet_utils import read_excel_all, write_excel_file

logger = logging.getLogger(__name__)

# Unsafe imports/operations to check for
UNSAFE_PATTERNS = [
    'import os', 'import sys', 'import subprocess', 'import shutil',
    '__import__', 'eval(', 'exec(', 'os.', 'sys.', 'subprocess.',
    'open(', 'file.', '.remove(', '.delete(', 'shutil.'
]

class ScriptExecutionService:
    def __init__(self):
        self.gemini = GeminiService()
        
        # Define safe operations (comprehensive set)
        self.safe_ops = {
            'read_file': self._safe_read_file,
            'write_file': self._safe_write_file,
            'read_excel': self._safe_read_excel,
            'write_excel': self._safe_write_excel,
            'read_excel_all': self._safe_read_excel_all,
            'write_excel_all': self._safe_write_excel_all,
            'read_json': self._safe_read_json,
            'write_json': self._safe_write_json,
            'read_image': self._safe_read_image,
            'write_image': self._safe_write_image,
            'read_docx': self._safe_read_docx,
            'write_docx': self._safe_write_docx,
            'encode_base64': self._safe_encode_base64,
            'decode_base64': self._safe_decode_base64
        }
        
        # Base modules that are always available
        self.allowed_modules = {
            # Data processing
            'pd': pd,
            'np': np,
            'json': json,
            
            # Document processing
            'Image': Image,
            'BytesIO': BytesIO,
            'base64': base64,
            
            # Excel Processing
            'Workbook': Workbook,
            'load_workbook': load_workbook,
            'get_column_letter': get_column_letter,
            'column_index_from_string': column_index_from_string,
            'Font': Font,
            'PatternFill': PatternFill, 
            'Alignment': Alignment,
            'Border': Border,
            'Side': Side,
            'NamedStyle': NamedStyle,
            'Color': Color,
            'Protection': Protection,
            'Fill': Fill,
            'COLOR_INDEX': COLOR_INDEX,
            'ColumnDimension': ColumnDimension,
            'RowDimension': RowDimension,
            'Cell': Cell,
            'XLImage': XLImage,
            'PyPDF2': PyPDF2,
            'Pt': Pt,
            'RGBColor': RGBColor,
            'Inches': Inches,
            'WD_ALIGN_PARAGRAPH': WD_ALIGN_PARAGRAPH,
        }
    
    def _safe_read_file(self, handle: str) -> bytes:
        """Safe file read operation"""
        return storage_service.download_file(handle)
    
    def _safe_write_file(self, filename: str, content: bytes) -> str:
        """Safe file write operation"""
        return storage_service.upload_file(filename, content)
    
    def _safe_read_excel(self, handle: str) -> pd.DataFrame:
        """Safe Excel read operation"""
        content = self._safe_read_file(handle)
        return pd.read_excel(BytesIO(content))
    
    def _safe_write_excel(self, df: pd.DataFrame, filename: str) -> str:
        """Safe Excel write operation"""
        buffer = BytesIO()
        df.to_excel(buffer)
        return self._safe_write_file(filename, buffer.getvalue())
    
    def _safe_read_excel_all(self, handle: str) -> Dict[str, pd.DataFrame]:
        """Read all sheets from Excel file"""
        content = self._safe_read_file(handle)
        return pd.read_excel(BytesIO(content), sheet_name=None)
    
    def _safe_write_excel_all(self, sheets: Dict[str, pd.DataFrame], filename: str) -> str:
        """Write multiple sheets to Excel file"""
        buffer = BytesIO()
        with pd.ExcelWriter(buffer) as writer:
            for sheet_name, df in sheets.items():
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        return self._safe_write_file(filename, buffer.getvalue())

    def _safe_read_json(self, handle: str) -> dict:
        """Safe JSON read operation"""
        content = self._safe_read_file(handle)
        return json.loads(content)

    def _safe_write_json(self, data: dict, filename: str) -> str:
        """Safe JSON write operation"""
        content = json.dumps(data).encode('utf-8')
        return self._safe_write_file(filename, content)

    def _safe_read_image(self, handle: str) -> Image:
        """Safe image read operation"""
        content = self._safe_read_file(handle)
        return Image.open(BytesIO(content))

    def _safe_write_image(self, img: Image, filename: str) -> str:
        """Safe image write operation"""
        buffer = BytesIO()
        img.save(buffer, format=img.format)
        return self._safe_write_file(filename, buffer.getvalue())

    def _safe_read_docx(self, handle: str) -> str:
        """Safe docx read operation"""
        content = self._safe_read_file(handle)
        doc = docx.Document(BytesIO(content))
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    
    def _safe_write_docx(self, text: str, filename: str) -> str:
        """Safe docx write operation"""
        doc = docx.Document()
        for paragraph in text.split('\n'):
            doc.add_paragraph(paragraph)
        buffer = BytesIO()
        doc.save(buffer)
        return self._safe_write_file(filename, buffer.getvalue())

    def _safe_encode_base64(self, data: bytes) -> str:
        """Safe base64 encode operation"""
        return base64.b64encode(data).decode('utf-8')
    
    def _safe_decode_base64(self, data: str) -> bytes:
        """Safe base64 decode operation"""
        return base64.b64decode(data)

    def generate_script(self, task: str, input_schema: Dict, output_requirements: Dict) -> str:
        """Generate clean Python code without markdown"""
        prompt = f"""Return ONLY Python code without markdown formatting.

                Task: {task}

                Available functions:
                Data Operations:
                - read_file(handle) -> bytes
                - write_file(filename, content) -> handle
                - read_json(handle) -> dict
                - write_json(data, filename) -> handle
                - encode_base64(data) -> str
                - decode_base64(data) -> bytes

                Excel Operations:
                - read_excel(handle) -> DataFrame
                - write_excel(df, filename) -> handle
                - read_excel_all(handle) -> Dict[str, DataFrame]
                - write_excel_all(sheets_dict, filename) -> handle

                Document Operations:
                - read_docx(handle) -> str
                - write_docx(text, filename) -> handle
                - read_image(handle) -> Image
                - write_image(img, filename) -> handle

                Available Modules:
                - pandas (pd)
                - numpy (np)
                - PIL.Image
                - json
                - openpyxl
                - docx

                {json.dumps(input_schema)}
                {json.dumps(output_requirements)}

                Notes:
                - DO NOT use file system operations
                - ONLY use provided functions
                - Include error handling
                - Return dictionary with results
                """
        result = self.gemini.analyze_text(prompt, "Generate Python code only")
        if result.get("success"):
            code = result["text"]
            if "```python" in code:
                code = code.split("```python")[1].split("```")[0]
            return code.strip()
        raise Exception(f"Script generation failed: {result.get('error')}")

    def _sanitize_script(self, script: str) -> str:
        """Sanitize script by checking for unsafe operations"""
        script_lines = script.split('\n')
        safe_lines = []
        
        for line in script_lines:
            line = line.strip()
            if any(pattern in line.lower() for pattern in UNSAFE_PATTERNS):
                continue  # Skip unsafe lines
            if line.startswith('import ') or line.startswith('from '):
                continue  # Skip all import statements
            safe_lines.append(line)
            
        return '\n'.join(safe_lines)

    def execute_script(self, script: str, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Execute a Python script with the provided input data"""
        try:
            # Remove any leading/trailing whitespace
            script = script.strip()
            
            # Remove common leading whitespace from every line
            lines = script.splitlines()
            if lines:
                # Find minimum indentation
                def get_indent(line):
                    return len(line) - len(line.lstrip()) if line.strip() else float('inf')
                min_indent = min(get_indent(line) for line in lines if line.strip())
                # Remove that amount of indentation from each line
                script = '\n'.join(
                    line[min_indent:] if line.strip() else ''
                    for line in lines
                )

            # Create a local namespace for script execution
            local_namespace = {
                'input_data': input_data,
                'read_excel_all': read_excel_all,
                'write_excel_all': write_excel_file,
                'get_excel_bytes': get_excel_bytes,
                'get_excel_bytes_from_dfs': get_excel_bytes_from_dfs,
                'pd': pd,
                'np': np
            }
            
            # Execute the script
            exec(script, globals(), local_namespace)
            
            # Get the output from local namespace
            output = local_namespace.get('output', {})
            
            # Check if output contains error
            if isinstance(output, dict) and 'error' in output:
                logger.error(f"Script execution produced error: {output['error']}")
                return {
                    'success': False,
                    'error': output['error']
                }
            
            return {
                'success': True,
                'output': output
            }
            
        except Exception as e:
            error_msg = f"Script execution failed: {str(e)}\nTraceback: {traceback.format_exc()}"
            logger.exception(error_msg)
            return {
                'success': False,
                'error': error_msg
            }
